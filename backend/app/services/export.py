"""导出服务：简历结构化内容导出为 Word（支持模板布局、主题色与排版密度）。"""

import io
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.resume import ResumeVersion

# 排版密度 → 正文基础字号（pt），与前端 DENSITIES 保持一致
DENSITY_FONT_SIZE = {"compact": 10.5, "standard": 11, "relaxed": 12}

# 需要双栏布局的模板
SIDEBAR_TEMPLATE = "sidebar"

_FALLBACK_COLOR = "1f2937"  # 无主题色时的标题色（深灰蓝）


def _normalize_hex(theme: str | None) -> str | None:
    """校验 #RRGGBB 主题色；非合法值（含历史 "blue" 等旧值）返回 None。"""
    if theme and re.fullmatch(r"#[0-9a-fA-F]{6}", theme):
        return theme[1:]
    return None


def _set_cell_shading(cell, hex_color: str) -> None:
    """设置单元格背景色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_paragraph_bottom_border(paragraph, hex_color: str) -> None:
    """设置段落底部边框（段落标题下划线）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _item_lines(item) -> list[str]:
    """条目 → 文本行列表（对象条目首键值对为首行，其余为后续行）。"""
    if isinstance(item, str):
        return [item]
    return [f"{k}: {v}" for k, v in item.items() if v]


def _styled_run(paragraph, text: str, *, bold: bool = False, size: float | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _render_single(doc: Document, info: dict, sections: list[dict], accent: str | None, name_fallback: str) -> None:
    """单栏布局：居中姓名 + 主题色段落标题（带下划线）。"""
    accent = accent or _FALLBACK_COLOR
    name = info.get("name") or name_fallback
    title = info.get("title", "")

    head = doc.add_heading(f"{name} {title}".strip(), level=0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in head.runs:
        run.font.color.rgb = RGBColor.from_string(accent)

    contact = info.get("contact", "")
    if info.get("location"):
        contact = f"{contact} · {info['location']}".strip()
    if contact:
        p = doc.add_paragraph(contact)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string("64748b")

    for section in sections:
        h = doc.add_heading(section.get("title", ""), level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(accent)
        _set_paragraph_bottom_border(h, accent)
        for item in section.get("items", []):
            lines = _item_lines(item)
            for i, line in enumerate(lines):
                p = doc.add_paragraph(line)
                if i == 0:
                    for run in p.runs:
                        run.bold = True


def _render_sidebar(doc: Document, info: dict, sections: list[dict], accent: str | None, base_size: float, name_fallback: str) -> None:
    """双栏侧栏布局：左列主题色底（基本信息 + 纯文本段落），右列其余段落。"""
    accent = accent or _FALLBACK_COLOR
    white = "FFFFFF"
    text_sections = [s for s in sections if all(isinstance(i, str) for i in s["items"])]
    left = text_sections
    right = [s for s in sections if s not in left]
    if not left and sections:
        left = [sections[0]]
        right = sections[1:]
    elif not right and sections:
        mid = (len(sections) + 1) // 2
        left = sections[:mid]
        right = sections[mid:]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Cm(5.4)
    right_cell.width = Cm(10.4)
    _set_cell_shading(left_cell, accent)

    # 左列：姓名 / 意向 / 联系方式 + 纯文本段落（白色）
    name = info.get("name") or name_fallback
    p = left_cell.paragraphs[0]
    _styled_run(p, name, bold=True, size=base_size + 3, color=white)
    if info.get("title"):
        _styled_run(left_cell.add_paragraph(), info["title"], size=base_size - 1, color=white)
    contact = info.get("contact", "")
    if info.get("location"):
        contact = f"{contact} · {info['location']}".strip()
    if contact:
        _styled_run(left_cell.add_paragraph(), contact, size=base_size - 1.5, color=white)
    for section in left:
        hp = left_cell.add_paragraph()
        _styled_run(hp, section.get("title", ""), bold=True, size=base_size, color=white)
        _set_paragraph_bottom_border(hp, white)
        for item in section.get("items", []):
            for line in _item_lines(item):
                _styled_run(left_cell.add_paragraph(), line, size=base_size - 1.5, color=white)

    # 右列：其余段落（主题色标题 + 下划线）
    for section in right:
        hp = right_cell.add_paragraph()
        _styled_run(hp, section.get("title", ""), bold=True, size=base_size + 1, color=accent)
        _set_paragraph_bottom_border(hp, accent)
        for item in section.get("items", []):
            lines = _item_lines(item)
            for i, line in enumerate(lines):
                p = right_cell.add_paragraph(line)
                if i == 0:
                    for run in p.runs:
                        run.bold = True


def resume_to_docx_bytes(version: ResumeVersion) -> bytes:
    """将结构化简历渲染为 Word 文档字节流（主题色 + 模板布局 + 密度字号）。"""
    content = version.content or {}
    accent = _normalize_hex(content.get("theme"))
    template = content.get("template", "classic")
    density = content.get("density", "standard")
    base_size = DENSITY_FONT_SIZE.get(density, 11)
    name_fallback = version.label or "简历"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(base_size)
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    info = content.get("info", {})
    sections = [s for s in content.get("sections", []) if s.get("title") and s.get("items")]

    if template == SIDEBAR_TEMPLATE:
        _render_sidebar(doc, info, sections, accent, base_size, name_fallback)
    else:
        _render_single(doc, info, sections, accent, name_fallback)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
